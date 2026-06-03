#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
apply_format.py —《新闻与传播研究》格式 skill 的「机械格式」处理器。

职责：只做非引文的版面格式调整，绝不改动任何文字内容、绝不改引文文本。
  - 页面：A4 + 页边距（每页约 38 行、每行 41 字）
  - 正文 / 各级标题 字体字号行距（参数全部来自 references/format_config.json，不写死）
  - 脚注字体 + 编号格式（①②③ 每页重排）
  - 摘要标签自动改为「内容提要」（仅标签）；英文标题/摘要缺失仅提醒、不自动生成
  - run() 返回变更数据（内存），不产生 JSON

所有参数从 references/format_config.json 读取；期刊改版只改该文件。
输出：<原名>_新传研究投稿版.docx（绝不覆盖原文件）。

用法（独立）：python3 apply_format.py <输入.docx> [-o 输出.docx]
常规请用：python3 apply_all.py <输入.docx>
"""

import os
import re
import sys
import json
import shutil
import argparse

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def load_config():
    cfg_path = os.path.join(
        os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
        "references", "format_config.json",
    )
    with open(cfg_path, encoding="utf-8") as f:
        return json.load(f)


CJK = re.compile(r"[一-鿿]")


def heading_level(text, cfg):
    t = text.strip()
    if not CJK.search(t) or len(t) > 40:
        return 0
    for lvl in ("h1", "h2", "h3", "h4"):
        pat = cfg["headings"].get(lvl, {}).get("prefix_regex")
        if pat and re.match(pat, t):
            return int(lvl[1])
    return 0


def set_run_font(run, eastasia, latin, size_pt=None, bold=None):
    r_pr = run._r.get_or_add_rPr()
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.font.bold = bold
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), latin)
    r_fonts.set(qn("w:hAnsi"), latin)
    r_fonts.set(qn("w:eastAsia"), eastasia)
    r_fonts.set(qn("w:cs"), latin)


def force_zero_spacing(p_elem):
    """直接在 w:spacing 上把段前/段后清零（含『行』单位 beforeLines/afterLines 与自动间距），
    避免 python-docx 对脚注段 space_before=0 不落属性、或残留 beforeLines/autospacing。
    不动 w:line / w:lineRule（行距由调用方另设）。"""
    pPr = p_elem.get_or_add_pPr()
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        pPr.append(spacing)
    for attr in ("before", "after", "beforeLines", "afterLines",
                 "beforeAutospacing", "afterAutospacing"):
        spacing.set(qn("w:" + attr), "0")


def apply_page(doc, cfg, changes):
    pg = cfg["page"]
    for section in doc.sections:
        section.page_width = Cm(pg["width_cm"])
        section.page_height = Cm(pg["height_cm"])
        section.top_margin = Cm(pg["margin_top_cm"])
        section.bottom_margin = Cm(pg["margin_bottom_cm"])
        section.left_margin = Cm(pg["margin_left_cm"])
        section.right_margin = Cm(pg["margin_right_cm"])
    changes.append("页面：A4，页边距 上下%scm/左右%scm（目标 每页38行×每行41字）"
                   % (pg["margin_top_cm"], pg["margin_left_cm"]))


def apply_body_and_headings(doc, cfg, changes):
    body = cfg["body"]
    n_body = 0
    n_head = {1: 0, 2: 0, 3: 0, 4: 0}
    for para in doc.paragraphs:
        if not para.text.strip():
            continue
        lvl = heading_level(para.text, cfg)
        if lvl == 0:
            for run in para.runs:
                set_run_font(run, body["font_eastasia"], body["font_latin"], body["size_pt"])
            pf = para.paragraph_format
            pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            pf.line_spacing = Pt(body["line_spacing_exact_pt"])
            pf.space_before = Pt(body.get("space_before_pt", 0))
            pf.space_after = Pt(body.get("space_after_pt", 0))
            force_zero_spacing(para._p)
            if pf.first_line_indent is None:
                pf.first_line_indent = Pt(body["size_pt"] * body["first_line_indent_chars"])
            n_body += 1
        else:
            h = cfg["headings"][f"h{lvl}"]
            for run in para.runs:
                set_run_font(run, h["font_eastasia"], body["font_latin"],
                             h["size_pt"], h.get("bold"))
            pf = para.paragraph_format
            pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            pf.line_spacing = Pt(body["line_spacing_exact_pt"])
            n_head[lvl] += 1
    changes.append("正文 %d 段：%s/%s %spt，固定行距 %spt，段前段后 0，首行缩进 %d 字"
                   % (n_body, body["font_eastasia"], body["font_latin"],
                      body["size_pt"], body["line_spacing_exact_pt"], body["first_line_indent_chars"]))
    changes.append("标题：H1×%d H2×%d H3×%d H4×%d（黑体，按级字号）"
                   % (n_head[1], n_head[2], n_head[3], n_head[4]))


def apply_footnote_numbering(doc, cfg, changes, flags):
    fn = cfg["footnote"]
    try:
        settings_el = doc.settings.element
        fn_pr = settings_el.find(qn("w:footnotePr"))
        if fn_pr is None:
            fn_pr = OxmlElement("w:footnotePr")
            settings_el.insert(0, fn_pr)
        # 先移除旧的 numFmt/numStart/numRestart
        for tag in (qn("w:numFmt"), qn("w:numStart"), qn("w:numRestart")):
            old = fn_pr.find(tag)
            if old is not None:
                fn_pr.remove(old)
        # OOXML schema 顺序：w:pos?, w:numFmt?, w:numStart?, w:numRestart?, w:footnote*
        # numFmt/numRestart 必须插在 w:footnote（分隔符引用 -1/0…）之前，否则 Word 报“不可读取的内容”并忽略编号
        sep_refs = fn_pr.findall(qn("w:footnote"))
        insert_idx = list(fn_pr).index(sep_refs[0]) if sep_refs else len(list(fn_pr))
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fn["num_fmt"])
        restart = OxmlElement("w:numRestart")
        restart.set(qn("w:val"), fn["num_restart"])
        fn_pr.insert(insert_idx, num_fmt)
        fn_pr.insert(insert_idx + 1, restart)
        changes.append("脚注编号：%s（①②③），%s（每页重排）"
                       % (fn["num_fmt"], fn["num_restart"]))
    except Exception as e:
        flags.append("脚注编号格式需在 Word/WPS 中手动设为 ①②③ 每页重排（自动设置失败：%s）" % e)


def apply_footnote_font(doc, cfg, changes):
    fn = cfg["footnote"]
    single = str(fn.get("line_spacing", "single")).lower() == "single"
    n = 0
    try:
        for footnote in doc.part.footnotes.footnotes:
            for para in footnote.paragraphs:
                # 脚注段落间距/行距由 footnote_ops.set_footnote_para_spacing 在写回时统一设置
                # （python-docx 对脚注段的 space_before=0 不稳定，改在 footnotes.xml 直接落属性）
                for run in para.runs:
                    set_run_font(run, fn["font_eastasia"], fn["font_latin"], fn["size_pt"])
                    n += 1
    except Exception:
        pass
    if n:
        changes.append("脚注字体：%s/%s %spt（小五），段前段后 0，%s"
                       % (fn["font_eastasia"], fn["font_latin"], fn["size_pt"],
                          "单倍行距" if single else "默认行距"))


ABS_LABEL_RE = re.compile(r"^(摘\s*要|提\s*要|内容提要)\s*")


def normalize_abstract_label(doc, cfg, changes, flags):
    """把摘要标签统一为本刊「内容提要」（仅改标签词，不动摘要正文）。"""
    target = cfg["front_matter"]["abstract_label"]
    for para in doc.paragraphs[:40]:
        t = para.text.strip()
        if not t:
            continue
        m = ABS_LABEL_RE.match(t)
        if not m:
            continue
        found = m.group(1).replace(" ", "")
        if found == target:
            return  # 已是「内容提要」
        # 在首个含标签的 run 内替换前缀（标签通常整段落在第一个文本 run）
        for run in para.runs:
            if not run.text:
                continue
            rm = ABS_LABEL_RE.match(run.text)
            if rm:
                run.text = target + run.text[rm.end():]
                changes.append("摘要标签「%s」→「%s」（仅标签，正文未动）" % (found, target))
                return
        flags.append("摘要标签为「%s」自动改写失败（标签跨多个 run），请手动改为「%s」。" % (found, target))
        return
    flags.append("未检出摘要/内容提要标签，请确认前置部分。")


def check_front_matter(doc, cfg, flags):
    """只做提醒，不自动补内容：本刊英文标题/摘要/关键词需作者另附（不能凭空翻译生成）。"""
    paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    has_en = any(len(re.findall(r"[A-Za-z]", p)) > 0.6 * len(p) and len(p) > 80
                 for p in paras)
    if not has_en:
        flags.append("未检出英文标题/摘要/关键词。本刊英文摘要在期刊末尾统一编排，"
                     "投稿通常需另附英文标题、作者、摘要、关键词——请作者自行补齐，"
                     "**本 skill 不自动翻译生成**（避免凭空创造内容）。")


def run(input_docx, output_docx=None, cfg=None):
    """执行机械格式，写出 docx，返回 {input, output, changes, flags}。不产 JSON。"""
    cfg = cfg or load_config()
    out = output_docx or (os.path.splitext(input_docx)[0] + "_新传研究投稿版.docx")
    shutil.copy2(input_docx, out)  # 绝不覆盖原文件
    doc = Document(out)
    changes, flags = [], []
    apply_page(doc, cfg, changes)
    apply_body_and_headings(doc, cfg, changes)
    apply_footnote_numbering(doc, cfg, changes, flags)
    apply_footnote_font(doc, cfg, changes)
    normalize_abstract_label(doc, cfg, changes, flags)
    check_front_matter(doc, cfg, flags)
    doc.save(out)
    return {"input": os.path.abspath(input_docx), "output": os.path.abspath(out),
            "format_changes": changes, "format_flags": flags}


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("docx")
    ap.add_argument("-o", "--out", default=None)
    args = ap.parse_args()
    if not os.path.exists(args.docx):
        print("文件不存在: %s" % args.docx, file=sys.stderr)
        sys.exit(1)
    r = run(args.docx, args.out)
    print("输出：%s" % r["output"])
    print("\n已完成的格式变更：")
    for c in r["format_changes"]:
        print("  ✓ " + c)
    if r["format_flags"]:
        print("\n待复核（格式类）：")
        for fl in r["format_flags"]:
            print("  ⚠ " + fl)


if __name__ == "__main__":
    main()
